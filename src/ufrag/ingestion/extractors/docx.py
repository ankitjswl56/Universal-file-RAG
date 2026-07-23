import docx

from ufrag.models import Section


def extract_docx(path) -> tuple[list[Section], list[str]]:
    document = docx.Document(str(path))
    paragraph_texts = [p.text for p in document.paragraphs]

    headings: list[tuple[int, str, int]] = []  # (level, title, paragraph_index)
    for i, para in enumerate(document.paragraphs):
        level = _heading_level(para.style.name if para.style else "")
        if level is not None and para.text.strip():
            headings.append((level, para.text.strip(), i))

    if not headings:
        content = "\n".join(paragraph_texts)
        sections = [
            Section(
                level=0,
                title="(untitled)",
                content=content,
                location={"para_start": 0, "para_end": max(len(paragraph_texts) - 1, 0)},
            )
        ]
        return sections, paragraph_texts

    sections = []
    for idx, (level, title, para_idx) in enumerate(headings):
        end_idx = headings[idx + 1][2] if idx + 1 < len(headings) else len(paragraph_texts)
        content = "\n".join(paragraph_texts[para_idx:end_idx])
        sections.append(
            Section(
                level=level,
                title=title,
                content=content,
                location={"para_start": para_idx, "para_end": end_idx - 1},
            )
        )
    return sections, paragraph_texts


def _heading_level(style_name: str) -> int | None:
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading "):
        suffix = style_name[len("Heading ") :]
        if suffix.isdigit():
            return int(suffix)
    return None
