"""Regenerates the DOCX and XLSX test fixtures under tests/fixtures/.

Run from the repo root: .venv/bin/python scripts/generate_docx_xlsx_fixtures.py
"""

import docx
import openpyxl

SECTIONS = [
    ("Overview", [
        "Continuum Task Tracker is a fictional project management tool used purely "
        "as a test fixture for the Universal File RAG DOCX ingestion pipeline. It "
        "is not a real product and none of the policies described here apply to "
        "any actual service.",
        "It organizes work into boards, lists, and cards, and this document "
        "describes plan limits, reminders, collaboration permissions, billing, "
        "and data export in enough detail to exercise heading-style-based "
        "chunking and paragraph citation accuracy across several sections.",
        "The document is intentionally verbose in places so its total length "
        "clears the length gate used to decide whether a hierarchical outline is "
        "worth building for a given file, rather than falling back to plain "
        "vector-only indexing.",
    ]),
    ("Plan Limits", [
        "Free plans may create up to 3 boards and 50 cards per board, with "
        "attachments capped at 10 MB per card. Cards beyond the 50 card limit "
        "cannot be created until existing cards are archived or deleted.",
        "Team plans remove the board and card count limits entirely and raise "
        "the per-card attachment cap to 250 MB. Team plans also unlock custom "
        "fields on cards, which Free plans cannot use under any circumstance.",
        "Exceeding the Free plan's attachment cap on a single card blocks new "
        "attachment uploads to that card specifically, but never blocks editing "
        "the card's text, checklist items, or due date.",
    ]),
    ("Reminders and Notifications", [
        "Due date reminders fire 24 hours and 1 hour before a card's due time by "
        "default, and this schedule can be customized per card on Team plans "
        "only; Free plan cards always use the default schedule.",
        "Notifications can be delivered by email, push notification, or both. "
        "Disabling all notification channels still allows due dates to display "
        "as overdue in red within the board itself, since that indicator is not "
        "considered a notification.",
        "Mentioning a teammate with @ in a card comment always sends an "
        "immediate notification regardless of that teammate's notification "
        "schedule preferences, since mentions are treated as high priority.",
    ]),
    ("Collaboration and Permissions", [
        "Team plan boards support three member roles: Viewer, Editor, and Admin. "
        "Viewers can read cards and comments but cannot modify or delete "
        "anything. Editors can create, edit, and move cards but cannot change "
        "board membership or delete the board itself.",
        "Admins have full control, including removing members and deleting the "
        "board. A board must always have at least one Admin; removing the last "
        "Admin is blocked until another member is promoted first.",
        "Free plan boards do not support roles at all: every member who has "
        "access can edit anything on the board, including deleting cards other "
        "members created.",
    ]),
    ("Billing and Refunds", [
        "Team plan subscriptions bill monthly or annually, chosen at signup, "
        "with annual billing discounted 15 percent versus paying monthly for "
        "twelve months. Billing occurs on the same calendar day each period as "
        "the original signup date.",
        "If a renewal payment fails, Continuum retries the charge every 3 days "
        "for up to 9 days before downgrading the workspace to Free. Downgrading "
        "never deletes boards or cards even if the workspace then exceeds Free "
        "plan limits; excess boards simply become read-only.",
        "Refunds are issued only for annual subscriptions cancelled within 14 "
        "days of the billing date, prorated to the unused portion of the year. "
        "Monthly subscriptions are never eligible for partial refunds.",
    ]),
    ("Data Export", [
        "Boards can be exported as CSV or JSON at any time from the board "
        "settings menu. CSV export flattens checklist items into a single "
        "semicolon-separated column, while JSON export preserves the full "
        "nested checklist structure.",
        "Workspace-wide export, which bundles every board into one archive, is "
        "a Team plan feature only; Free plan workspaces can only export one "
        "board at a time.",
        "Exported files never include attachment binaries themselves, only "
        "attachment filenames and their original upload URLs, to keep export "
        "archive sizes manageable.",
    ]),
    ("Account Deletion", [
        "When a user deletes their Continuum account entirely, all boards they "
        "solely own are scheduled for permanent deletion after a 30 day grace "
        "period, during which the account can be reactivated by signing back in.",
        "Boards with other Admins are not deleted when one Admin's account is "
        "removed; ownership simply passes to the remaining Admins, and the "
        "board continues to function normally for every other member.",
        "After the 30 day grace period ends, deletion is irreversible and "
        "Continuum cannot recover any content under any circumstance, even upon "
        "a direct request from the former account owner.",
    ]),
    ("Integrations", [
        "Team plans support two-way sync with a small number of external "
        "calendar providers; due dates set on a card automatically create a "
        "matching calendar event, and moving a card's due date updates that "
        "event within a few minutes.",
        "Free plans support one-way export to calendar only: a card's due date "
        "can be added as a calendar event manually via an export link, but "
        "changes made afterward in Continuum do not propagate to that event.",
        "Disconnecting an integration does not delete calendar events already "
        "created by it; those events remain in the external calendar until the "
        "user removes them manually, since Continuum has no way to reach back "
        "into a calendar it is no longer connected to.",
    ]),
    ("Support Response Times", [
        "Free plan support requests are answered over email only, with a "
        "target response time of 5 business days and no guaranteed resolution "
        "time frame attached to that target.",
        "Team plan support requests are answered within 1 business day, and "
        "workspaces with 10 or more paid seats additionally get access to live "
        "chat support during business hours in the workspace's configured "
        "timezone, excluding weekends and public holidays observed there.",
    ]),
    ("Security", [
        "Team plans support single sign-on via SAML for the whole workspace, "
        "which when enabled requires every member to authenticate through the "
        "configured identity provider and disables password-based login for "
        "that workspace entirely.",
        "Free plans do not support single sign-on and always use password-based "
        "login, optionally combined with two-factor authentication enabled "
        "individually by each member from their own account settings.",
        "Audit logs recording membership changes, permission changes, and board "
        "deletions are retained for 90 days on Team plans and are not available "
        "at all on Free plans.",
    ]),
]


def make_structured_docx(path: str) -> None:
    document = docx.Document()
    for title, paragraphs in SECTIONS:
        document.add_heading(title, level=1)
        for para in paragraphs:
            document.add_paragraph(para)
    document.save(path)


def make_short_docx(path: str) -> None:
    document = docx.Document()
    document.add_paragraph(
        "Reminder: the office guest wifi password is Sunflower42, rotated monthly. "
        "Ask Priya at the front desk for parking validation stickers."
    )
    document.save(path)


def make_budget_xlsx(path: str) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Q1 Budget"

    sheet.append(["Category", "Budgeted", "Actual", "Variance"])
    rows = [
        ("Engineering", 50000, 47500),
        ("Marketing", 20000, 23000),
        ("Design", 15000, 14200),
        ("Operations", 10000, 9800),
    ]
    for i, (category, budgeted, actual) in enumerate(rows, start=2):
        sheet.cell(row=i, column=1, value=category)
        sheet.cell(row=i, column=2, value=budgeted)
        sheet.cell(row=i, column=3, value=actual)
        sheet.cell(row=i, column=4, value=f"=C{i}-B{i}")

    total_row = len(rows) + 2
    sheet.cell(row=total_row, column=1, value="Total")
    sheet.cell(row=total_row, column=2, value=f"=SUM(B2:B{total_row - 1})")
    sheet.cell(row=total_row, column=3, value=f"=SUM(C2:C{total_row - 1})")
    sheet.cell(row=total_row, column=4, value=f"=SUM(D2:D{total_row - 1})")

    workbook.save(path)


if __name__ == "__main__":
    make_structured_docx("tests/fixtures/sample_manual.docx")
    make_short_docx("tests/fixtures/sample_note.docx")
    make_budget_xlsx("tests/fixtures/sample_budget.xlsx")
    print("done")
